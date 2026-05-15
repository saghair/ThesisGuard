from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING


@dataclass
class ParagraphInfo:
    text: str
    style_name: str
    font_names: list[str] = field(default_factory=list)
    font_sizes_pt: list[float] = field(default_factory=list)
    line_spacings: list[float] = field(default_factory=list)


@dataclass
class ParsedDocument:
    full_text: str
    sections: list[str]
    paragraphs: list[ParagraphInfo]
    heading_texts: list[str]
    page_margins_cm: dict[str, float] | None
    file_type: str = "docx"


EMU_PER_CM = 360000
TWIPS_PER_LINE = 240
PDF_HEADING_SIZE_THRESHOLD = 13.0


def emu_to_cm(value: int | None) -> float | None:
    if value is None:
        return None
    return round(value / EMU_PER_CM, 2)


def _extract_line_spacing(paragraph) -> list[float]:
    fmt = paragraph.paragraph_format
    rule = fmt.line_spacing_rule
    spacing = fmt.line_spacing

    if spacing is None:
        return []
    if rule == WD_LINE_SPACING.MULTIPLE:
        if isinstance(spacing, float):
            return [round(spacing, 2)]
        if isinstance(spacing, int):
            return [round(spacing / TWIPS_PER_LINE, 2)]
    if rule in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST):
        return []
    if isinstance(spacing, (int, float)) and 0.5 <= float(spacing) <= 5.0:
        return [round(float(spacing), 2)]
    return []


def parse_docx(path: Path) -> ParsedDocument:
    document = Document(str(path))
    paragraphs: list[ParagraphInfo] = []
    heading_texts: list[str] = []
    section_titles: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else "Unknown"
        font_names: list[str] = []
        font_sizes: list[float] = []

        for run in paragraph.runs:
            if run.font.name:
                font_names.append(run.font.name)
            if run.font.size:
                try:
                    font_sizes.append(round(run.font.size.pt, 1))
                except Exception:
                    pass

        info = ParagraphInfo(
            text=text,
            style_name=style_name,
            font_names=font_names,
            font_sizes_pt=font_sizes,
            line_spacings=_extract_line_spacing(paragraph),
        )
        paragraphs.append(info)

        if text and style_name.lower().startswith("heading"):
            heading_texts.append(text)
            section_titles.append(text)

    margins = None
    if document.sections:
        first = document.sections[0]
        margins = {
            "top": emu_to_cm(first.top_margin),
            "bottom": emu_to_cm(first.bottom_margin),
            "left": emu_to_cm(first.left_margin),
            "right": emu_to_cm(first.right_margin),
        }
        margins = {k: v for k, v in margins.items() if v is not None}

    full_text = "\n".join(p.text for p in paragraphs if p.text)
    return ParsedDocument(
        full_text=full_text,
        sections=section_titles,
        paragraphs=paragraphs,
        heading_texts=heading_texts,
        page_margins_cm=margins if margins else None,
        file_type="docx",
    )


def parse_pdf(path: Path) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber is not installed. Run: pip install pdfplumber")

    paragraphs: list[ParagraphInfo] = []
    heading_texts: list[str] = []
    section_titles: list[str] = []
    page_margins_cm: dict[str, float] | None = None

    with pdfplumber.open(str(path)) as pdf:
        if pdf.pages:
            first_page = pdf.pages[0]
            pw = first_page.width
            ph = first_page.height
            cb = first_page.cropbox
            if cb:
                pts_per_cm = 28.3465
                page_margins_cm = {
                    "left": round(cb[0] / pts_per_cm, 2),
                    "top": round(cb[1] / pts_per_cm, 2),
                    "right": round((pw - cb[2]) / pts_per_cm, 2),
                    "bottom": round((ph - cb[3]) / pts_per_cm, 2),
                }

        for page in pdf.pages:
            words = page.extract_words(extra_attrs=["fontname", "size"])
            if not words:
                continue

            lines: dict[float, list[dict]] = {}
            for word in words:
                top = round(float(word.get("top", 0)), 1)
                lines.setdefault(top, []).append(word)

            for top_y in sorted(lines.keys()):
                line_words = lines[top_y]
                line_text = " ".join(w["text"] for w in line_words).strip()
                if not line_text:
                    continue

                font_sizes = [
                    round(float(w["size"]), 1)
                    for w in line_words if w.get("size")
                ]
                font_names = list({
                    w["fontname"] for w in line_words if w.get("fontname")
                })

                avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 0.0
                is_heading = avg_size >= PDF_HEADING_SIZE_THRESHOLD and len(line_text.split()) <= 12

                info = ParagraphInfo(
                    text=line_text,
                    style_name="Heading" if is_heading else "Normal",
                    font_names=font_names,
                    font_sizes_pt=font_sizes,
                    line_spacings=[],
                )
                paragraphs.append(info)

                if is_heading:
                    heading_texts.append(line_text)
                    section_titles.append(line_text)

    full_text = "\n".join(p.text for p in paragraphs if p.text)
    return ParsedDocument(
        full_text=full_text,
        sections=section_titles,
        paragraphs=paragraphs,
        heading_texts=heading_texts,
        page_margins_cm=page_margins_cm,
        file_type="pdf",
    )


def parse_document(path: Path) -> ParsedDocument:
    """Auto-detect file type and parse accordingly."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    elif suffix == ".docx":
        return parse_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
