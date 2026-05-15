from __future__ import annotations
import re
from pathlib import Path


def extract_requirements_from_guidance(path: Path, template_name: str) -> dict:
    """
    Read a thesis guidance document and extract content requirements.
    Works with both heading-based and table-based documents.
    """
    suffix = path.suffix.lower()

    # ── Read full text including tables ──────────────────────────────
    full_text = ""
    heading_texts = []

    if suffix == ".docx":
        from docx import Document
        doc = Document(str(path))
        # Paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                full_text += p.text + "\n"
                if p.style and p.style.name.lower().startswith("heading"):
                    heading_texts.append(p.text.strip())
        # Tables (important — this doc is table-heavy)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text += cell.text + "\n"
    elif suffix == ".pdf":
        from app.services.parser import parse_pdf
        parsed = parse_pdf(path)
        full_text = parsed.full_text
        heading_texts = parsed.heading_texts
    else:
        raise ValueError(f"Unsupported file: {suffix}")

    lower = full_text.lower()

    # ── 1. Required sections ──────────────────────────────────────────
    KNOWN_SECTIONS = [
        "Abstract", "Introduction", "Literature Review", "Research Review",
        "Theoretical Framework", "Methodology", "Methods", "Analysis",
        "Original Part", "Project Part", "Implementation", "Results",
        "Discussion", "General Conclusion", "Conclusion", "References",
        "Bibliography", "Appendix", "Topicality", "Problem Formulation",
    ]

    # Detect chapter titles from text patterns
    chapter_matches = re.findall(
        r'chapter\s+\d+\s*[–\-—|]\s*([A-Za-z][^\n\|]{3,50})',
        full_text, re.IGNORECASE
    )
    chapter_sections = []
    for m in chapter_matches:
        clean = m.strip().strip('*').strip()
        if clean and len(clean) < 60:
            chapter_sections.append(clean)

    required_sections = []
    for sec in KNOWN_SECTIONS:
        if sec.lower() in lower:
            required_sections.append(sec)
    for ch in chapter_sections[:6]:
        if ch not in required_sections:
            required_sections.append(ch)

    # Deduplicate
    seen, unique_sections = set(), []
    for s in required_sections:
        if s.lower() not in seen:
            seen.add(s.lower())
            unique_sections.append(s)

    # ── 2. Heading patterns ───────────────────────────────────────────
    heading_patterns = []
    for i in range(1, 5):
        if f"chapter {i}" in lower:
            heading_patterns.append(f"Chapter {i}")
    if "general conclusion" in lower:
        heading_patterns.append("General Conclusion")

    # ── 3. Citation style ─────────────────────────────────────────────
    citation_style = None
    for style in ["IEEE", "APA", "MLA", "Harvard", "Chicago"]:
        if style.lower() in lower:
            citation_style = style
            break

    # ── 4. Extract specific content rules ────────────────────────────
    rules = []

    if "smart" in lower:
        rules.append("Research aim must be SMART (Specific, Measurable, Achievable, Relevant, Time-bound)")
    if "defended thesis" in lower or "nail 1" in lower or "three nails" in lower:
        rules.append("Must include Defended Thesis — 3 statements: 'Has been offered / modified / suggested'")
    if "every chapter must end" in lower or "chapter must end" in lower:
        rules.append("Every chapter must end with a short Conclusion paragraph")
    if "general conclusion" in lower:
        rules.append("Must include a General Conclusion (~1 page) summarising all findings")
    if "object of research" in lower:
        rules.append("Must clearly define the Object of Research (system/company/process being studied)")
    if "subject of research" in lower:
        rules.append("Must clearly define the Subject of Research (specific aspect analysed)")
    if "topicality" in lower:
        rules.append("Topicality section required — fact-based, 3–4 sentences with statistics")
    if "problem formulation" in lower or "research problem" in lower:
        rules.append("Problem Formulation required — 3–4 concrete sentences citing facts")
    if "research aim" in lower or "research goal" in lower:
        rules.append("Research Aim (Goal) required — one measurable, specific sentence")
    if "planned results" in lower:
        rules.append("Planned Results required — specific and measurable, state who benefits")
    if "bi-weekly" in lower or "biweekly" in lower:
        rules.append("Bi-weekly progress emails required to supervisor")
    if "annotated bibliography" in lower:
        rules.append("Annotated bibliography required (minimum 10 sources)")

    # Page length rules
    page_matches = re.findall(r'(\d+)\s*[–\-—]\s*(\d+)\s*pages?', full_text, re.IGNORECASE)
    for mn, mx in page_matches[:4]:
        rules.append(f"Expected length: {mn}–{mx} pages for this section")

    # Source requirements
    source_matches = re.findall(r'min(?:imum)?\.?\s*(\d+)\s*(?:academic\s+)?sources?', full_text, re.IGNORECASE)
    if source_matches:
        rules.append(f"Minimum {source_matches[0]} academic sources required in references")

    notes = (
        f"Requirements extracted from '{path.name}'. "
        + (" | ".join(rules) if rules else "Upload detected no specific content rules.")
    )

    config = {
        "name": template_name,
        "document_type": "thesis",
        "allowed_fonts": ["Times New Roman"],
        "allowed_font_sizes": [12],
        "line_spacing": 1.5,
        "margins_cm": {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.5},
        "required_sections": unique_sections[:15],
        "heading_patterns": heading_patterns,
        "citation_style": citation_style,
        "notes": notes,
    }

    return config, rules
